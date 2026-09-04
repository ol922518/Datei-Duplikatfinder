"""
document_viewer.py
-------------------
Eingebauter Datei-Viewer für die Ergebnis-Tabelle: zeigt die zur aktuell
ausgewählten Zeile gehörende Datei an, damit man ihren Inhalt vor dem
Verschieben prüfen kann - ursprünglich 1:1 vom Datei-Umbenenner übernommen,
seitdem um Touchpad-Gesten ergänzt (siehe unten) - der engine-Import zeigt
hier auf duplicate_engine statt rename_engine.

Unterstützt:
- PDF (natives QtPdf-Widget)
- Bilder (JPEG, PNG, TIFF, HEIC, ...) über Qts eingebaute Bildformate
- Text/Markdown/CSV/JSON/YAML (reiner Text, Markdown mit einfacher Formatierung)
- Word (.docx) als reiner Text (über python-docx, sofern installiert)

Bedienung bei Bildern/PDF: Zwei-Finger-Wischen auf dem Trackpad scrollt
(hoch/runter/links/rechts, über Qts eingebaute Scroll-Behandlung von
QScrollArea/QPdfView), Zusammen-/Auseinanderziehen (Pinch) zoomt (siehe
eventFilter() unten, macOS-Trackpad-Geste). Zoom UND Scroll-Position
bleiben dabei beim Wechsel zur nächsten Datei erhalten (als relativer
Bruchteil 0..1 je Achse, nicht als Pixelwert - siehe _scroll_fraction) -
war z.B. die rechte untere Ecke der vorherigen Datei zu sehen, zeigt die
nächste Datei ebenfalls ihre rechte untere Ecke, unabhängig von deren
tatsächlicher Größe. Der allererste Start (noch nie manuell gezoomt/
gescrollt) ist echte 100% oben links; "↺ Einpassen" setzt jederzeit
bewusst darauf zurück.

Alles läuft lokal, ohne Internetzugriff - passend zum Rest der App. Einzige
Ausnahme: bei echten Kamerafotos mit GPS-Daten zeigt eine kleine Metadaten-
Zeile Datum/Kamera/Koordinaten an, mit einem Button "🌐 Ort ermitteln", der
bewusst erst auf Klick eine Online-Abfrage bei OpenStreetMap auslöst (siehe
duplicate_engine.reverse_geocode()) - nie automatisch.
"""

from pathlib import Path

from PySide6.QtCore import QEvent, Qt
from PySide6.QtGui import QPixmap
from PySide6.QtPdf import QPdfDocument
from PySide6.QtPdfWidgets import QPdfView
from PySide6.QtWidgets import (
    QApplication,
    QHBoxLayout,
    QLabel,
    QPushButton,
    QScrollArea,
    QStackedWidget,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

import duplicate_engine as engine

try:
    import docx  # python-docx - Verfügbarkeit über engine.HAS_DOCX geprüft
    # (dieselbe Bibliothek, dort bereits erkannt), hier nur zusätzlich für
    # den direkten docx.Document(...)-Aufruf unten importiert.
except ImportError:
    pass
MARKDOWN_EXTENSIONS = {".md"}

ZOOM_STEP = 1.25
ZOOM_MIN = 0.1
ZOOM_MAX = 5.0


class DocumentViewer(QWidget):
    """Zeigt eine einzelne Datei passend zu ihrem Typ an. Aufruf über
    `show_file(path)`, `clear()` setzt die Anzeige zurück."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self._current_pixmap: QPixmap | None = None
        # None = automatisch auf Fensterbreite einpassen; sonst manuell
        # gewählter Zoomfaktor (1.0 = Originalgröße), nur für Bilder - für
        # PDF hält QPdfView seinen Zoom-Zustand selbst (zoomMode/zoomFactor).
        # Bleibt (wie der Zoom-Zustand von self.pdf_view) bewusst über einen
        # Dateiwechsel hinweg erhalten (siehe show_file()). Ausgangswert
        # echte 1.0 (100%) statt None/Einpassen, damit die allererste
        # angezeigte Datei ebenfalls bei echten 100% startet.
        self._image_zoom: float | None = 1.0
        # Relative Scroll-Position (0.0..1.0 je Achse) der zuletzt aktiven
        # Bild-/PDF-Ansicht - bewusst als Bruchteil statt Pixelwert, damit
        # z.B. "rechts unten" auch bei unterschiedlich großen Dateien
        # vergleichbar bleibt. Wird bei jedem manuellen Scrollen aktualisiert
        # (siehe _on_scroll_changed()) und beim Öffnen der nächsten Datei
        # wiederhergestellt (siehe _restore_scroll_position()), damit sich
        # beim Vergleichen mehrerer Duplikate dieselbe Stelle im Blick
        # behalten lässt.
        self._scroll_fraction: tuple[float, float] = (0.0, 0.0)
        # GPS-Koordinaten des aktuell angezeigten Fotos (falls vorhanden) -
        # Grundlage für den Button "🌐 Ort ermitteln" (siehe _on_geocode_clicked).
        self._current_photo_gps: tuple[float, float] | None = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        self.title_label = QLabel("Vorschau")
        self.title_label.setStyleSheet("font-weight: bold;")
        self.title_label.setWordWrap(True)
        layout.addWidget(self.title_label)

        # Zoom-Leiste - nur bei Bildern/PDF sinnvoll, wird je nach
        # angezeigtem Inhalt aktiviert/deaktiviert (siehe _set_active_page).
        self.zoom_bar = QWidget()
        zoom_layout = QHBoxLayout(self.zoom_bar)
        zoom_layout.setContentsMargins(0, 0, 0, 4)
        zoom_out_btn = QPushButton("➖")
        zoom_out_btn.setFixedWidth(32)
        zoom_out_btn.setToolTip("Verkleinert die Vorschau.")
        zoom_out_btn.clicked.connect(self._zoom_out)
        zoom_layout.addWidget(zoom_out_btn)
        self.zoom_label = QLabel("100%")
        self.zoom_label.setAlignment(Qt.AlignCenter)
        self.zoom_label.setFixedWidth(48)
        zoom_layout.addWidget(self.zoom_label)
        zoom_in_btn = QPushButton("➕")
        zoom_in_btn.setFixedWidth(32)
        zoom_in_btn.setToolTip("Vergrößert die Vorschau.")
        zoom_in_btn.clicked.connect(self._zoom_in)
        zoom_layout.addWidget(zoom_in_btn)
        zoom_fit_btn = QPushButton("↺ Einpassen")
        zoom_fit_btn.setToolTip("Setzt den Zoom zurück, sodass die Vorschau wieder in den verfügbaren Platz passt.")
        zoom_fit_btn.clicked.connect(self._zoom_fit)
        zoom_layout.addWidget(zoom_fit_btn)
        zoom_layout.addStretch(1)
        layout.addWidget(self.zoom_bar)

        # Metadaten-Zeile - nur bei echten Kamerafotos (EXIF-Datum/GPS)
        # sichtbar, siehe _update_photo_meta_bar().
        self.photo_meta_bar = QWidget()
        photo_meta_layout = QHBoxLayout(self.photo_meta_bar)
        photo_meta_layout.setContentsMargins(0, 0, 0, 4)
        self.photo_meta_label = QLabel()
        self.photo_meta_label.setWordWrap(True)
        self.photo_meta_label.setStyleSheet("color: palette(mid);")
        photo_meta_layout.addWidget(self.photo_meta_label, 1)
        self.geocode_btn = QPushButton("🌐 Ort ermitteln")
        self.geocode_btn.setToolTip(
            "Fragt den Ortsnamen zu den GPS-Koordinaten dieses Fotos online bei "
            "OpenStreetMap ab (einzige Stelle in der App, die dafür Internet braucht - "
            "geschieht nur auf diesen Klick hin, nie automatisch). Danach über 'Vorschau "
            "aktualisieren' als Baustein {ort} nutzbar."
        )
        self.geocode_btn.clicked.connect(self._on_geocode_clicked)
        photo_meta_layout.addWidget(self.geocode_btn)
        self.photo_meta_bar.setVisible(False)
        layout.addWidget(self.photo_meta_bar)

        self.stack = QStackedWidget()
        layout.addWidget(self.stack, 1)

        # Leer-Seite: solange (noch) nichts ausgewählt ist.
        self.empty_page = QLabel("Zeile in der Tabelle auswählen, um eine Vorschau zu sehen.")
        self.empty_page.setAlignment(Qt.AlignCenter)
        self.empty_page.setWordWrap(True)
        self.empty_page.setStyleSheet("color: palette(mid);")
        self.stack.addWidget(self.empty_page)

        # Hinweis-Seite: Datei ohne Vorschau bzw. Lesefehler.
        self.unsupported_page = QLabel()
        self.unsupported_page.setAlignment(Qt.AlignCenter)
        self.unsupported_page.setWordWrap(True)
        self.unsupported_page.setStyleSheet("color: palette(mid);")
        self.stack.addWidget(self.unsupported_page)

        # Text/Markdown/Word-Ansicht.
        self.text_view = QTextEdit()
        self.text_view.setReadOnly(True)
        self.stack.addWidget(self.text_view)

        # Bild-Ansicht.
        self.image_scroll = QScrollArea()
        self.image_scroll.setWidgetResizable(True)
        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_scroll.setWidget(self.image_label)
        self.stack.addWidget(self.image_scroll)

        # PDF-Ansicht. Ausgangswert Custom/100% statt FitToWidth (siehe
        # show_file()) - FitToWidth berechnet je nach Seiten-/Fenstergröße
        # auch mal einen Wert über 100%, was beim allerersten Öffnen einer
        # PDF-Datei überraschend/zufällig wirken würde.
        self.pdf_document = QPdfDocument(self)
        self.pdf_view = QPdfView()
        self.pdf_view.setDocument(self.pdf_document)
        self.pdf_view.setPageMode(QPdfView.PageMode.MultiPage)
        self.pdf_view.setZoomMode(QPdfView.ZoomMode.Custom)
        self.pdf_view.setZoomFactor(1.0)
        self.stack.addWidget(self.pdf_view)

        # Pinch-Zoom per Trackpad (macOS: "Zusammen-/Auseinanderziehen") -
        # Qt liefert das als natives Gesten-Event an das Widget unter dem
        # Mauszeiger, hier also die Viewports von Bild-/PDF-Ansicht (siehe
        # eventFilter() unten). Zwei-Finger-Scrollen (Pan) braucht dagegen
        # keinen eigenen Code - QScrollArea/QPdfView verarbeiten normale
        # Trackpad-Scroll-Gesten bereits eingebaut als Wheel-Events; über die
        # Scrollbar-Signale wird dieselbe Bewegung zusätzlich als relative
        # Position gemerkt (siehe _on_scroll_changed()).
        self.image_scroll.viewport().installEventFilter(self)
        self.pdf_view.viewport().installEventFilter(self)
        self.image_scroll.horizontalScrollBar().valueChanged.connect(self._on_scroll_changed)
        self.image_scroll.verticalScrollBar().valueChanged.connect(self._on_scroll_changed)
        self.pdf_view.horizontalScrollBar().valueChanged.connect(self._on_scroll_changed)
        self.pdf_view.verticalScrollBar().valueChanged.connect(self._on_scroll_changed)

        self._set_active_page(self.empty_page)

    def eventFilter(self, obj, event) -> bool:
        if event.type() == QEvent.NativeGesture and event.gestureType() == Qt.NativeGestureType.ZoomNativeGesture:
            # event.value() ist der kleine Skalierungs-Zuwachs seit dem
            # letzten Gesten-Frame (z.B. 0.02) - direkt als Faktor auf den
            # aktuellen Zoom anwendbar.
            self._adjust_zoom(1 + event.value())
            return True
        return super().eventFilter(obj, event)

    def clear(self) -> None:
        self.title_label.setText("Vorschau")
        self._current_pixmap = None
        self._update_photo_meta_bar(None)
        self._set_active_page(self.empty_page)

    def show_file(self, path: Path) -> None:
        self.title_label.setText(path.name)
        self._current_pixmap = None
        # Zoom (self._image_zoom bzw. der Zoom-Zustand von self.pdf_view)
        # UND Scroll-Position (self._scroll_fraction) werden bewusst NICHT
        # zurückgesetzt - bleiben über den Dateiwechsel hinweg erhalten
        # (Ausgangswert beider ist echte 100% oben links, siehe __init__),
        # bis der Nutzer selbst "↺ Einpassen" klickt. So bleibt z.B. eine in
        # die rechte untere Ecke gezoomte Stelle beim Vergleichen mehrerer
        # Duplikate an derselben Stelle sichtbar.
        self._update_photo_meta_bar(None)

        if not path.exists():
            self._show_message(f"Datei nicht gefunden:\n{path.name}")
            return

        ext = path.suffix.lower()
        if ext == ".pdf":
            self._show_pdf(path)
        elif ext == ".docx":
            self._show_docx(path)
        elif ext in engine.TEXT_EXTENSIONS:
            self._show_text(path, markdown=ext in MARKDOWN_EXTENSIONS)
        else:
            # Alles andere (auch unbekannte Endungen) als Bild versuchen -
            # Qt kennt deutlich mehr Bildformate, als die App fürs {content}
            # aktiv unterstützt (z.B. auch HEIC).
            self._show_image_or_unsupported(path)

    # ------------------------------------------------------------------
    def _set_active_page(self, widget) -> None:
        """Wechselt die angezeigte Seite und schaltet die Zoom-Leiste passend
        dazu ein (Bild/PDF) oder aus (Text/Hinweis/leer)."""
        self.stack.setCurrentWidget(widget)
        self.zoom_bar.setEnabled(widget in (self.image_scroll, self.pdf_view))
        self._update_zoom_label()

    def _show_message(self, text: str) -> None:
        self.unsupported_page.setText(text)
        self._set_active_page(self.unsupported_page)

    def _show_pdf(self, path: Path) -> None:
        # load() liefert einen QPdfDocument.Error-Wert (nicht Status - ein
        # Vergleich mit Status.Ready wäre deshalb immer ungleich und somit
        # wirkungslos), erfolgreich ist genau Error.None_.
        error = self.pdf_document.load(str(path))
        if error != QPdfDocument.Error.None_:
            self._show_message(f"PDF konnte nicht geöffnet werden:\n{path.name}")
            return
        # Zoom-Zustand NICHT zurückgesetzt (siehe show_file()) - Ausgangswert
        # Custom/100% kommt bereits aus __init__.
        self._restore_scroll_position(self.pdf_view)
        self._set_active_page(self.pdf_view)

    def _show_image_or_unsupported(self, path: Path) -> None:
        pixmap = QPixmap(str(path))
        if pixmap.isNull():
            self._show_message(f"Keine Vorschau verfügbar für:\n{path.name}")
            return
        self._current_pixmap = pixmap
        self._apply_scaled_pixmap()
        self._restore_scroll_position(self.image_scroll)
        self._update_photo_meta_bar(engine.extract_photo_metadata(path))
        self._set_active_page(self.image_scroll)

    @staticmethod
    def _format_place(lat: float, lon: float, cached: dict | None) -> str:
        """Baut die Anzeige '📍 Ort, Land (Koordinaten)' - Ort/Land nur, wenn
        bereits per reverse_geocode() ermittelt (siehe LOCATION_CACHE), sonst
        nur die reinen Koordinaten."""
        label = ""
        if cached:
            label = cached.get("ort", "")
            if cached.get("land"):
                label += (", " if label else "") + cached["land"]
        return f"📍 {label} ({lat:.5f}, {lon:.5f})" if label else f"📍 {lat:.5f}, {lon:.5f}"

    def _update_photo_meta_bar(self, meta: dict | None) -> None:
        """Zeigt bei echten Kamerafotos (EXIF-Datum und/oder GPS vorhanden)
        eine kleine Metadaten-Zeile mit Button '🌐 Ort ermitteln' an - sonst
        (Screenshots/Grafiken ohne EXIF, oder andere Dateitypen) bleibt sie
        ausgeblendet."""
        self._current_photo_gps = None
        if not meta or not (meta.get("date") or meta.get("latitude") is not None):
            self.photo_meta_bar.setVisible(False)
            return

        parts = []
        if meta.get("date"):
            parts.append(f"📅 {meta['date']}")
        if meta.get("camera"):
            parts.append(f"📷 {meta['camera']}")
        lat, lon = meta.get("latitude"), meta.get("longitude")
        if lat is not None and lon is not None:
            self._current_photo_gps = (lat, lon)
            cached = engine.LOCATION_CACHE.get(engine.location_cache_key(lat, lon))
            parts.append(self._format_place(lat, lon, cached))
        self.photo_meta_label.setText("   ".join(parts))
        self.geocode_btn.setVisible(self._current_photo_gps is not None)
        self.geocode_btn.setEnabled(True)
        self.geocode_btn.setText("🌐 Ort ermitteln")
        self.photo_meta_bar.setVisible(True)

    def _on_geocode_clicked(self) -> None:
        if self._current_photo_gps is None:
            return
        lat, lon = self._current_photo_gps
        self.geocode_btn.setEnabled(False)
        self.geocode_btn.setText("Suche…")
        # Sorgt dafür, dass der deaktivierte Button/Text vor dem (kurz
        # blockierenden) Netzwerkaufruf gleich sichtbar wird.
        QApplication.processEvents()
        result = engine.reverse_geocode(lat, lon)
        # None = Abfrage selbst fehlgeschlagen (z.B. kein Internet); ein
        # Ergebnis-dict mit leeren Werten ist dagegen eine erfolgreiche
        # Abfrage, die für diese Koordinate nur nichts gefunden hat - beides
        # braucht eine unterschiedliche Rückmeldung.
        if result is None:
            self.geocode_btn.setText("⚠ Kein Internet?")
        elif result.get("ort") or result.get("land"):
            self.photo_meta_label.setText(
                self.photo_meta_label.text().rsplit("📍", 1)[0] + self._format_place(lat, lon, result)
            )
            self.geocode_btn.setText("🌐 Ort ermitteln")
        else:
            self.geocode_btn.setText("⚠ Kein Ort gefunden")
        self.geocode_btn.setEnabled(True)

    def _current_image_zoom(self) -> float:
        """Aktueller Zoomfaktor fürs Bild (1.0 = Originalgröße) - im
        Einpassen-Modus (self._image_zoom is None) aus der verfügbaren
        Breite berechnet, sonst der manuell gewählte Wert."""
        if self._current_pixmap is None or self._current_pixmap.width() <= 0:
            return 1.0
        if self._image_zoom is not None:
            return self._image_zoom
        viewport_width = max(10, self.image_scroll.viewport().width() - 4)
        return min(1.0, viewport_width / self._current_pixmap.width())

    def _apply_scaled_pixmap(self) -> None:
        if self._current_pixmap is None:
            return
        zoom = self._current_image_zoom()
        width = max(10, round(self._current_pixmap.width() * zoom))
        self.image_label.setPixmap(self._current_pixmap.scaledToWidth(width, Qt.SmoothTransformation))
        self._update_zoom_label()

    def _show_text(self, path: Path, markdown: bool) -> None:
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
        except OSError as e:
            self._show_message(f"Datei konnte nicht gelesen werden:\n{e}")
            return
        if markdown:
            self.text_view.setMarkdown(text)
        else:
            self.text_view.setPlainText(text)
        self._set_active_page(self.text_view)

    def _show_docx(self, path: Path) -> None:
        if not engine.HAS_DOCX:
            self._show_message(
                "Keine Vorschau möglich: Paket 'python-docx' ist nicht installiert."
            )
            return
        try:
            document = docx.Document(str(path))
            text = "\n\n".join(p.text for p in document.paragraphs if p.text.strip())
        except Exception as e:
            self._show_message(f"Word-Datei konnte nicht gelesen werden:\n{e}")
            return
        self.text_view.setPlainText(text or "(leeres Dokument)")
        self._set_active_page(self.text_view)

    # ------------------------------------------------------------------
    # Zoom (Bilder + PDF)
    # ------------------------------------------------------------------
    def _zoom_in(self) -> None:
        self._adjust_zoom(ZOOM_STEP)

    def _zoom_out(self) -> None:
        self._adjust_zoom(1 / ZOOM_STEP)

    def _adjust_zoom(self, factor: float) -> None:
        current = self.stack.currentWidget()
        if current is self.image_scroll and self._current_pixmap is not None:
            self._image_zoom = max(ZOOM_MIN, min(ZOOM_MAX, self._current_image_zoom() * factor))
            self._apply_scaled_pixmap()
        elif current is self.pdf_view:
            new_factor = max(ZOOM_MIN, min(ZOOM_MAX, self.pdf_view.zoomFactor() * factor))
            self.pdf_view.setZoomMode(QPdfView.ZoomMode.Custom)
            self.pdf_view.setZoomFactor(new_factor)
            self._update_zoom_label()

    def _zoom_fit(self) -> None:
        current = self.stack.currentWidget()
        self._scroll_fraction = (0.0, 0.0)
        if current is self.image_scroll:
            self._image_zoom = None
            self._apply_scaled_pixmap()
        elif current is self.pdf_view:
            self.pdf_view.setZoomMode(QPdfView.ZoomMode.FitToWidth)
            self._update_zoom_label()

    def _on_scroll_changed(self, _value: int = 0) -> None:
        """Merkt sich die aktuelle Scroll-Position der gerade aktiven Bild-/
        PDF-Ansicht als Bruchteil (0.0..1.0) je Achse statt als Pixelwert -
        so bleibt z.B. "rechts unten" auch bei unterschiedlich großen
        Dateien vergleichbar (siehe _restore_scroll_position())."""
        current = self.stack.currentWidget()
        if current is self.image_scroll:
            view = self.image_scroll
        elif current is self.pdf_view:
            view = self.pdf_view
        else:
            return
        hbar, vbar = view.horizontalScrollBar(), view.verticalScrollBar()
        self._scroll_fraction = (
            hbar.value() / hbar.maximum() if hbar.maximum() else 0.0,
            vbar.value() / vbar.maximum() if vbar.maximum() else 0.0,
        )

    def _restore_scroll_position(self, view) -> None:
        """Stellt die zuletzt gemerkte relative Scroll-Position
        (self._scroll_fraction) in der übergebenen Bild-/PDF-Ansicht
        wieder her - aufgerufen beim Öffnen einer neuen Datei, damit eine
        z.B. in die rechte untere Ecke gezoomte Stelle beim Wechsel
        erhalten bleibt."""
        hbar, vbar = view.horizontalScrollBar(), view.verticalScrollBar()
        x_frac, y_frac = self._scroll_fraction
        hbar.setValue(round(x_frac * hbar.maximum()))
        vbar.setValue(round(y_frac * vbar.maximum()))

    def _update_zoom_label(self) -> None:
        # Hinweis: QPdfView.zoomFactor() liefert im Modus "FitToWidth" nicht
        # den tatsächlich angezeigten Skalierungsfaktor (bleibt am zuletzt
        # manuell gesetzten Wert stehen) - deshalb dort "Auto" statt einer
        # (falschen) Prozentzahl anzeigen.
        current = self.stack.currentWidget()
        if current is self.image_scroll and self._current_pixmap is not None:
            if self._image_zoom is None:
                self.zoom_label.setText("Auto")
            else:
                self.zoom_label.setText(f"{round(self._image_zoom * 100)}%")
        elif current is self.pdf_view:
            if self.pdf_view.zoomMode() == QPdfView.ZoomMode.FitToWidth:
                self.zoom_label.setText("Auto")
            else:
                self.zoom_label.setText(f"{round(self.pdf_view.zoomFactor() * 100)}%")
        else:
            self.zoom_label.setText("–")

    # ------------------------------------------------------------------
    def resizeEvent(self, event) -> None:
        super().resizeEvent(event)
        if self._current_pixmap is not None and self.stack.currentWidget() is self.image_scroll:
            self._apply_scaled_pixmap()
